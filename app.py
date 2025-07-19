import os
import requests
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Setup Flask App
app = Flask(__name__)
# This allows your Render frontend to talk to your Render backend
CORS(app) 

# --- AI Formatting Endpoint ---
@app.route('/format-resume-ai', methods=['POST'])
def format_resume_with_ai():
    try:
        data = request.get_json()
        if not data or 'cleanedText' not in data:
            return jsonify({"error": "Missing cleanedText in request"}), 400
            
        cleaned_text = data['cleanedText']

        # Gets the API Key from the Environment Variable you will set on Render
        api_key = os.environ.get('GEMINI_API_KEY')
        if not api_key:
            return jsonify({"error": "API key is not configured on the server."}), 500

        # The full, working prompt
        prompt = f"""
You are an expert-level resume parser and formatter. Your task is to take pre-cleaned text and produce a clean, ATS-friendly, single-column resume. You must follow a strict, multi-step process.

**Step 1: Analyze and Reconstruct**
First, analyze the text. It may be jumbled from a multi-column layout. Identify all the logical sections (Name, Contact, Summary, Experience, Skills, Education, etc.) and reconstruct them into a single, ordered column. Your highest priority is correctly grouping all content under its proper heading.

**Step 2: Apply Section-Specific Bullet Point Rules**
After reconstructing the content, you MUST apply these specific rules for creating bullet points. This is the most critical step.

* **Rule A: For the "PROFESSIONAL EXPERIENCE" section:**
    * If a job description contains lines that explicitly start with a bullet character ('•', '*', '-'), EACH of those lines MUST become a separate bullet point ('•') in the output.
    * If a job description contains multiple distinct paragraphs that do NOT start with a bullet character, EACH of those paragraphs MUST also become a separate bullet point ('•').
    * Merge any wrapped lines that belong to the same paragraph or bullet point.

* **Rule B: For the "TECHNICAL SKILLS" or "SKILLS" section:**
    * This section has a different format. If you see sub-categories (e.g., "OPERATING SYSTEMS"), these sub-categories are headings.
    * The list of skills that follows a sub-category (e.g., "Linux, Windows") should be a **single line of plain text**.
    * **You MUST NOT add a '•' bullet point to the sub-category heading OR to the list of skills in this section.**

* **Rule C: For all other sections:**
    * Only create a bullet point if the line explicitly started with a bullet character in the original text. Otherwise, treat it as a normal paragraph.

**Step 3: Final Formatting**
Finally, apply these overall formatting rules to the structured text from Step 2.

* **Headers:** All main section headers (PROFESSIONAL EXPERIENCE, TECHNICAL SKILLS, etc.) must be in ALL CAPS.
* **Contact Info:** The line directly under the candidate's name should contain all contact details, separated by " | ".
* **Cleanup:** Preserve all URLs. All other junk symbols should already have been removed.

**OUTPUT REQUIREMENTS:**
* Produce ONLY the final, clean, formatted resume text.
* Do not include any explanations, notes, or apologies.
* The output must begin directly with the candidate's name.

[START OF CLEANED RESUME TEXT]
{cleaned_text}
[END OF CLEANED RESUME TEXT]
"""
        
        api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"
        
        payload = {
            "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            "generationConfig": {
                "temperature": 0.0,
                "topK": 1
            }
        }
        
        response = requests.post(api_url, json=payload)
        response.raise_for_status()
        
        result = response.json()

        if result.get("candidates") and result["candidates"][0].get("content", {}).get("parts", [{}])[0].get("text"):
            formatted_text = result["candidates"][0]["content"]["parts"][0]["text"]
            return jsonify({"formattedText": formatted_text})
        else:
            print("API Error Response:", result)
            return jsonify({"error": "Failed to get valid response from AI model."}), 500

    except requests.exceptions.RequestException as e:
        print(f"Error calling Gemini API: {e}")
        return jsonify({"error": f"Failed to communicate with AI service: {e}"}), 500
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return jsonify({"error": "An internal server error occurred."}), 500


# --- DOCX Creation Endpoint ---
@app.route('/create-resume', methods=['POST'])
def create_resume_docx():
    try:
        data = request.get_json()
        if not data or 'resumeText' not in data:
            return jsonify({"error": "Missing resumeText in request"}), 400

        plain_text = data['resumeText']
        lines = plain_text.strip().split('\n')

        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        for i, line in enumerate(lines):
            trimmed_line = line.strip()
            if not trimmed_line:
                continue 

            if i == 0:
                p = document.add_paragraph()
                run = p.add_run(trimmed_line)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                continue

            if i == 1:
                p = document.add_paragraph(trimmed_line)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                continue
            
            if (len(trimmed_line) > 2 and trimmed_line == trimmed_line.upper() and 
                '•' not in trimmed_line and any(c.isalpha() for c in trimmed_line)):
                p = document.add_paragraph()
                run = p.add_run(trimmed_line)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                continue
                
            if trimmed_line.startswith('•'):
                bullet_text = trimmed_line.lstrip('• ').strip()
                p = document.add_paragraph(bullet_text, style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Inches(0.25)
                continue
            
            p = document.add_paragraph(trimmed_line)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Error creating document: {e}")
        return jsonify({"error": "Failed to create .docx file on server"}), 500
